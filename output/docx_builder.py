"""
Word document builder using python-docx.
Generates formatted .docx files from Estimate and letter text.
"""
import io
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from models.estimate_models import Estimate, ProductEstimate, LineItem


def _set_font(run, bold=False, size=11, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _set_font(run, bold=True, size=14 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_company_header(doc: Document):
    company = os.getenv("COMPANY_NAME", "DRX Roofing & General Contracting")
    license_num = os.getenv("COMPANY_LICENSE", "")
    phone = os.getenv("COMPANY_PHONE", "")
    email = os.getenv("COMPANY_EMAIL", "")
    address = os.getenv("COMPANY_ADDRESS", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company)
    _set_font(run, bold=True, size=16, color=(31, 73, 125))

    for detail in [address, phone, email, license_num]:
        if detail:
            p = doc.add_paragraph(detail)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _set_font(run, size=10)

    doc.add_paragraph()  # spacer


def build_estimate_docx(estimate: Estimate) -> bytes:
    """Build a formatted estimate .docx and return as bytes."""
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    _add_company_header(doc)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ROOFING MATERIAL ESTIMATE")
    _set_font(run, bold=True, size=14)

    # Job info
    doc.add_paragraph()
    info_lines = [
        ("Date:", estimate.date_prepared or str(date.today())),
        ("Property:", estimate.property_address or "Subject Property"),
        ("Prepared By:", estimate.prepared_by),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label} ")
        _set_font(run_label, bold=True, size=11)
        run_value = p.add_run(value)
        _set_font(run_value, size=11)

    if estimate.narrative:
        doc.add_paragraph()
        p = doc.add_paragraph(estimate.narrative)
        p.paragraph_format.space_after = Pt(12)

    # Per-product tables
    for pe in estimate.products:
        doc.add_paragraph()
        _add_heading(doc, f"{pe.manufacturer} {pe.product_line}" + (f" — {pe.color}" if pe.color else ""), level=2)

        # Line item table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        headers = ["Category", "Description", "Qty", "Unit", "Total"]
        header_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                _set_font(run, bold=True, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for li in pe.line_items:
            row = table.add_row()
            cells = row.cells
            cells[0].text = li.category
            cells[1].text = li.description
            cells[2].text = f"{li.quantity:.2f}"
            cells[3].text = li.unit
            cells[4].text = f"${li.total:,.2f}"
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    _set_font(run, size=10)

        # Subtotal row
        subtotal_row = table.add_row()
        subtotal_row.cells[0].merge(subtotal_row.cells[3])
        subtotal_row.cells[0].text = "MATERIAL SUBTOTAL"
        for run in subtotal_row.cells[0].paragraphs[0].runs:
            _set_font(run, bold=True, size=10)
        subtotal_row.cells[4].text = f"${pe.subtotal_materials:,.2f}"
        subtotal_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in subtotal_row.cells[4].paragraphs[0].runs:
            _set_font(run, bold=True, size=10)

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph(estimate.disclaimer)
    p.paragraph_format.space_before = Pt(12)
    for run in p.runs:
        _set_font(run, size=9, color=(128, 128, 128))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_letter_docx(letter_text: str, title: str = "Correspondence") -> bytes:
    """Build a formatted letter .docx from plain text and return as bytes."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    _add_company_header(doc)

    # Split letter into paragraphs and add
    for para_text in letter_text.split("\n"):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            _set_font(run, size=11)
        p.paragraph_format.space_after = Pt(0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
